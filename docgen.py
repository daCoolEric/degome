"""
Study guide (markdown) -> branded Word document (.docx), via python-docx.

Kept deliberately dependency-light so it runs in the slim cloud image.
Handles the structures the guide prompt produces: ## / ### headings,
bullet and numbered lists, **bold**, *italic*, `code`, and paragraphs.
"""

import re
from datetime import date
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

TEAL = RGBColor(0x05, 0x63, 0x6D)
TEAL_DARK = RGBColor(0x03, 0x39, 0x40)
GOLD = RGBColor(0xB8, 0x86, 0x1F)   # print-friendly gold
GREY = RGBColor(0x59, 0x59, 0x59)

_INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def _add_runs(paragraph, text: str, base_size=11):
    """Render **bold**, *italic* and `code` inside one paragraph."""
    for token in _INLINE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(base_size - 1)
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            paragraph.add_run(token)


def _is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and s.count("|") >= 2


def _is_separator_row(line: str) -> bool:
    s = line.strip().strip("|")
    return bool(s) and all(ch in " :-|" for ch in s) and "-" in s


def _cells(line: str):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _add_table(doc, rows):
    """Emit a proper Word table from collected markdown pipe rows."""
    header = None
    body = rows
    if len(rows) >= 2 and _is_separator_row(rows[1]):
        header = _cells(rows[0])
        body = [r for r in rows[2:] if not _is_separator_row(r)]
    else:
        body = [r for r in rows if not _is_separator_row(r)]

    data = ([header] if header else []) + [_cells(r) for r in body]
    if not data:
        return
    cols = max(len(r) for r in data)
    table = doc.add_table(rows=len(data), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(data):
        for c in range(cols):
            cell = table.cell(i, c)
            cell.text = ""
            p = cell.paragraphs[0]
            _add_runs(p, row[c] if c < len(row) else "", base_size=10)
            for run in p.runs:
                run.font.size = run.font.size or Pt(10)
                if header and i == 0:
                    run.bold = True
                    run.font.color.rgb = TEAL
    doc.add_paragraph()


def guide_to_docx(markdown: str, lecture_name: str) -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---- title block ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Degome Study Guide")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = TEAL

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"{lecture_name}  \u00b7  {date.today().strftime('%d %B %Y')}")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = GREY

    doc.add_paragraph()

    # ---- body ----
    in_code = False
    table_rows = []

    def flush_table():
        nonlocal table_rows
        if table_rows:
            _add_table(doc, table_rows)
            table_rows = []

    for raw in markdown.splitlines():
        line = raw.rstrip()

        if not in_code and _is_table_row(line):
            table_rows.append(line)
            continue
        flush_table()

        if line.strip().startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            p = doc.add_paragraph()
            r = p.add_run(raw)
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
            continue

        if not line.strip():
            continue

        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            p = doc.add_heading(level=min(level, 3))
            p.text = ""
            r = p.add_run(m.group(2).strip())
            if level <= 2:
                r.font.size = Pt(15)
                r.font.color.rgb = TEAL
            else:
                r.font.size = Pt(12.5)
                r.font.color.rgb = TEAL_DARK
            r.font.bold = True
            continue

        m = re.match(r"^\s*[-*]\s+(.*)", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, m.group(1))
            continue

        m = re.match(r"^\s*\d+[.)]\s+(.*)", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, m.group(1))
            continue

        p = doc.add_paragraph()
        _add_runs(p, line)

    flush_table()

    # ---- footer line ----
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = foot.add_run("\u2014  Nunya, adidoe; asi metun\u025b o  \u2014")
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = GOLD

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
